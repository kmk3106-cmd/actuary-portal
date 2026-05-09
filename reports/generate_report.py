# -*- coding: utf-8 -*-
"""
IFRS17 계리결산 보고서 자동생성 스크립트
대상 파일: (이동민)IFRS17재무결산분석_202603_월별.xlsx
출력 파일: Actuarial Report_202603.docx
"""

import openpyxl
from openpyxl.utils import column_index_from_string
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime

# ──────────────────────────────────────────────────────────────────────────────
# 경로 설정
# ──────────────────────────────────────────────────────────────────────────────
BASE_DIR   = r"C:\Users\USER\actuary potal\reports"
INPUT_FILE = os.path.join(BASE_DIR, "input",
                          "(이동민)IFRS17재무결산분석_202603_월별.xlsx")
OUTPUT_FILE = os.path.join(BASE_DIR, "output", "Actuarial Report_202603.docx")

# ──────────────────────────────────────────────────────────────────────────────
# 유틸리티 함수
# ──────────────────────────────────────────────────────────────────────────────
def fmt_amt(val, unit=100_000_000, decimals=1):
    """원 → 억원 변환"""
    if val is None or val == '':
        return '-'
    try:
        v = float(val) / unit
        fmt = f"{{:,.{decimals}f}}"
        return fmt.format(v)
    except Exception:
        return str(val)


def fmt_eok(val, decimals=1):
    """억원 단위 값 포맷 (이미 억원인 경우)"""
    if val is None or val == '':
        return '-'
    try:
        v = float(val)
        fmt = f"{{:,.{decimals}f}}"
        return fmt.format(v)
    except Exception:
        return str(val)


def safe(val):
    """None → [데이터 없음]"""
    if val is None:
        return '[데이터 없음]'
    return val


# ──────────────────────────────────────────────────────────────────────────────
# 스타일 헬퍼
# ──────────────────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    """셀 배경색 설정"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """셀 테두리 설정"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom),
                      ('left', left), ('right', right)]:
        if val is not None:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val)
            el.set(qn('w:sz'), '4')
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), '000000')
            tcBorders.append(el)
    tcPr.append(tcBorders)


def style_header_cell(cell, text, font_size=10, bold=True,
                      bg_color='1F497D', font_color='FFFFFF',
                      align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = text
    cell.paragraphs[0].alignment = align
    run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else \
          cell.paragraphs[0].add_run(text)
    if not cell.paragraphs[0].runs:
        cell.text = ''
        run = cell.paragraphs[0].add_run(text)
    else:
        run = cell.paragraphs[0].runs[0]
        run.text = text
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(
        int(font_color[0:2], 16),
        int(font_color[2:4], 16),
        int(font_color[4:6], 16)
    )
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_bg(cell, bg_color)


def write_cell(cell, text, font_size=9, bold=False,
               align=WD_ALIGN_PARAGRAPH.RIGHT, indent=False):
    cell.text = ''
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(str(text))
    run.font.size = Pt(font_size)
    run.font.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if indent:
        para.paragraph_format.left_indent = Cm(0.3)


def add_section_title(doc, text, level=1):
    """섹션 제목 단락 추가"""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    run.font.size = Pt(13) if level == 1 else Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return para


def add_subtitle(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(text)
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return para


# ──────────────────────────────────────────────────────────────────────────────
# 데이터 읽기
# ──────────────────────────────────────────────────────────────────────────────
print("엑셀 파일 로딩 중...")
wb = openpyxl.load_workbook(INPUT_FILE, data_only=True)

# --- Setting_Report ---
ws_set = wb['Setting_Report']
report_ym = ws_set.cell(row=4, column=3).value   # 202603 (기시)
if report_ym is None:
    report_ym = ws_set.cell(row=5, column=3).value
print(f"  결산기준월: {report_ym}")

# --- 손익_요약 ---
ws_pl = wb['손익_요약']
title_cell = ws_pl.cell(row=2, column=2).value   # ▣ 2026.03月 결산보고

# 열 인덱스 매핑 (억원 시트 기준으로 읽음)
COL_BC = column_index_from_string('BC')   # 55 → 2026.03 당월
COL_BD = column_index_from_string('BD')   # 56 → 전월比
COL_BE = column_index_from_string('BE')   # 57 → 3개월 누계
COL_BF = column_index_from_string('BF')   # 58 → 월평균

# 손익_요약 실제 행 번호 (탐색 결과 기준)
ROW_MAP = {
    'ins_pl':        7,   # Ⅰ 보험손익
    'ins_rev':       8,   # Ⅰ.1 보험수익
    'csm_amort':     9,   # CSM 상각
    'ra_change':    10,   # RA 변동
    'ins_claim':    11,   # 보험금(예상-실제)
    'acq_cost':     12,   # 신계약비(예상-실제)
    'maint_cost':   13,   # 유지비
    'inv_mgmt':     14,   # 투자관리비
    'loss_adj':     15,   # 손해조사비
    'acf_alloc':    16,   # 보험취득CF배분
    'lob_alloc':    17,   # 손실부담비용 배분
    'other':        18,   # 기타
    'ins_svc_exp':  19,   # Ⅰ.2 보험서비스비용
    'occ_claim':    20,   # 발생보험금
    'maint_act':    21,   # 유지비(실제)
    'inv_act':      22,   # 투자관리비(실제)
    'ladj_act':     23,   # 손해조사비(실제)
    'loss_lob':     24,   # 손실부담비용 전/환입
    'ins_fin_pl':   27,   # Ⅱ 보험금융손익
    'ins_fin_rev':  28,   # Ⅱ.1 보험금융수익
    'ins_fin_exp':  35,   # Ⅱ.2 보험금융비용
    'oci':          42,   # Ⅲ 기타포괄손익
}

def read_pl_row(row_key):
    """손익_요약 시트에서 당월/전월比/누계 읽기 (억원 단위)"""
    r = ROW_MAP[row_key]
    cur  = ws_pl.cell(row=r, column=COL_BC).value
    mom  = ws_pl.cell(row=r, column=COL_BD).value
    cum3 = ws_pl.cell(row=r, column=COL_BE).value
    return cur, mom, cum3

def get_label(row_key):
    b = ws_pl.cell(row=ROW_MAP[row_key], column=2).value
    c = ws_pl.cell(row=ROW_MAP[row_key], column=3).value
    return b or c or row_key

# 예실차 (BJ~BM)
COL_BJ = column_index_from_string('BJ')  # 62
COL_BK = column_index_from_string('BK')  # 63
COL_BL = column_index_from_string('BL')  # 64
COL_BM = column_index_from_string('BM')  # 65

var_rows = []
for r in range(7, 13):  # 보험금~계
    label = ws_pl.cell(row=r, column=COL_BJ).value
    exp   = ws_pl.cell(row=r, column=COL_BK).value
    act   = ws_pl.cell(row=r, column=COL_BL).value
    diff  = ws_pl.cell(row=r, column=COL_BM).value
    if label is not None:
        var_rows.append((label, exp, act, diff))

# 최근 12개월 추세 (AR~BC, 열44~55)
TREND_COLS = list(range(44, 56))  # AR~BC
trend_headers = [ws_pl.cell(row=6, column=c).value for c in TREND_COLS]

def read_trend_row(row_key):
    r = ROW_MAP[row_key]
    return [ws_pl.cell(row=r, column=c).value for c in TREND_COLS]

trend_ins_pl   = read_trend_row('ins_pl')
trend_fin_pl   = read_trend_row('ins_fin_pl')
trend_oci      = read_trend_row('oci')

# --- 손익_억 (기초/기말 잔액) ---
ws_eok = wb['손익_억']
# 잔여 기말 Row43, 기초 Row44 (D=구분, E=BEL, F=RA, G=CSM, H=LOSS, I=OCI)
def read_balance(row):
    bel  = ws_eok.cell(row=row, column=5).value
    ra   = ws_eok.cell(row=row, column=6).value
    csm  = ws_eok.cell(row=row, column=7).value
    loss = ws_eok.cell(row=row, column=8).value
    oci  = ws_eok.cell(row=row, column=9).value
    return bel, ra, csm, loss, oci

end_r = read_balance(43)   # 잔여 기말
beg_r = read_balance(44)   # 잔여 기초
end_d = (ws_eok.cell(row=74, column=5).value,
         ws_eok.cell(row=74, column=6).value,
         None,
         None,
         ws_eok.cell(row=74, column=9).value)
beg_d = (ws_eok.cell(row=75, column=5).value,
         ws_eok.cell(row=75, column=6).value,
         None,
         None,
         ws_eok.cell(row=75, column=9).value)

def add_bal(a, b):
    """None safe 덧셈"""
    try:
        return (a or 0) + (b or 0)
    except Exception:
        return None

# 합산 기말/기초 (발생사고부채 포함)
bal_end = [add_bal(end_r[i], end_d[i]) for i in range(5)]
bal_beg = [add_bal(beg_r[i], beg_d[i]) for i in range(5)]
bal_chg = [add_bal(bal_end[i], -(bal_beg[i] or 0)) if bal_end[i] is not None else None
           for i in range(5)]

# --- 이행현금흐름상세 (가정변경) ---
ws_cf = wb['이행현금흐름상세']
cf_code_labels = {
    '006': '가정변경 (할인율/공시이율)',
    '151': '모델변경',
    '152': '계리가정변경 (사업비율)',
    '153': '계리가정변경 (위험률)',
    '154': '계리가정변경 (해지율)',
    '155': '계리가정변경 (기타)',
    '185': 'VFA 기업의 몫 조정',
}

# BEL, RA, CSM, LOSS 변동 섹션별로 읽기
cf_sections = {
    'BEL': (6, 14),     # 행 7~14
    'RA':  (15, 23),    # 행 16~23
    'CSM': (24, 32),    # 행 25~32
    'LOSS': (33, 41),   # 행 34~41
}

cf_data = {}
for section, (row_start_marker, row_end_marker) in cf_sections.items():
    rows = []
    for r in range(row_start_marker + 1, row_end_marker + 1):
        code = ws_cf.cell(row=r, column=1).value
        if code is not None and str(code) in cf_code_labels:
            v202601 = ws_cf.cell(row=r, column=2).value
            v202602 = ws_cf.cell(row=r, column=3).value
            v202603 = ws_cf.cell(row=r, column=4).value
            rows.append((str(code), cf_code_labels[str(code)], v202601, v202602, v202603))
    cf_data[section] = rows

print("  데이터 읽기 완료")
print(f"  손익_요약 제목: {title_cell}")
print(f"  Ⅰ 보험손익 당월: {read_pl_row('ins_pl')[0]:.1f}억원")
print(f"  BEL 기말: {bal_end[0]:,.1f}억원")

# ──────────────────────────────────────────────────────────────────────────────
# Word 문서 생성
# ──────────────────────────────────────────────────────────────────────────────
print("Word 문서 생성 중...")
doc = Document()

# 페이지 여백 설정
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.0)

# 기본 폰트 설정
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(9)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# ══════════════════════════════════════════════════════════════════════════════
# [제목 페이지]
# ══════════════════════════════════════════════════════════════════════════════
# 상단 여백
for _ in range(6):
    doc.add_paragraph()

# 제목
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run("IFRS17 계리결산 보고서")
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
run.font.name = '맑은 고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

doc.add_paragraph()

# 부제
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run("2026년 3월 결산  (기준월: 202603)")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
run.font.name = '맑은 고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

for _ in range(4):
    doc.add_paragraph()

# 구분선
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run("─" * 40)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

doc.add_paragraph()

# 메타 정보
for label, value in [("작  성", "계리결산팀"), ("일  자", "2026.03.31"),
                      ("기준월", str(report_ym))]:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"{label} :  {value}")
    run.font.size = Pt(12)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# Section 1. 손익 요약
# ══════════════════════════════════════════════════════════════════════════════
add_section_title(doc, "Section 1.  손익 요약")
para = doc.add_paragraph()
run = para.add_run("(단위: 억원)")
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

# 표 정의
pl_rows = [
    # (들여쓰기 여부, 키, 표시 레이블, 굵게)
    (False, 'ins_pl',       'Ⅰ  보험손익',              True),
    (False, 'ins_rev',      'Ⅰ.1  보험수익',             True),
    (True,  'csm_amort',    '  - CSM 상각',               False),
    (True,  'ra_change',    '  - RA 변동',                False),
    (True,  'ins_claim',    '  - 보험금 (예상-실제)',      False),
    (True,  'acq_cost',     '  - 신계약비 (예상-실제)',    False),
    (True,  'maint_cost',   '  - 유지비',                  False),
    (True,  'inv_mgmt',     '  - 투자관리비',              False),
    (True,  'loss_adj',     '  - 손해조사비',              False),
    (True,  'acf_alloc',    '  - 보험취득CF배분',          False),
    (True,  'lob_alloc',    '  - 손실부담비용 배분',       False),
    (False, 'ins_svc_exp',  'Ⅰ.2  보험서비스비용',        True),
    (True,  'occ_claim',    '  - 발생보험금',              False),
    (True,  'maint_act',    '  - 유지비 (실제)',            False),
    (True,  'inv_act',      '  - 투자관리비 (실제)',        False),
    (True,  'ladj_act',     '  - 손해조사비 (실제)',        False),
    (True,  'loss_lob',     '  - 손실부담비용 전/환입',    False),
    (False, 'ins_fin_pl',   'Ⅱ  보험금융손익',            True),
    (False, 'ins_fin_rev',  'Ⅱ.1  보험금융수익',          True),
    (False, 'ins_fin_exp',  'Ⅱ.2  보험금융비용',          True),
    (False, 'oci',          'Ⅲ  기타포괄손익',            True),
]

tbl = doc.add_table(rows=len(pl_rows) + 1, cols=4)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

# 열 너비
col_widths = [Cm(7.0), Cm(3.5), Cm(3.5), Cm(4.0)]
for i, w in enumerate(col_widths):
    for cell in tbl.columns[i].cells:
        cell.width = w

# 헤더
hdrs = ['항목', '당월(억원)', '전월比(억원)', '3개월 누계(억원)']
hdr_colors = ['1F497D', '2E74B5', '2E74B5', '2E74B5']
for ci, (h, c) in enumerate(zip(hdrs, hdr_colors)):
    style_header_cell(tbl.rows[0].cells[ci], h, bg_color=c)

# 데이터 행
for ri, (indent, key, label, bold) in enumerate(pl_rows, start=1):
    cur, mom, cum3 = read_pl_row(key)
    row_cells = tbl.rows[ri].cells

    # 항목명
    write_cell(row_cells[0], label, bold=bold,
               align=WD_ALIGN_PARAGRAPH.LEFT)
    if bold:
        set_cell_bg(row_cells[0], 'DDEEFF')

    # 숫자
    for ci, val in enumerate([cur, mom, cum3], start=1):
        write_cell(row_cells[ci], fmt_eok(val) if val is not None else '-',
                   bold=bold)
        if bold:
            set_cell_bg(row_cells[ci], 'DDEEFF')

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# Section 2. 보험부채 잔액
# ══════════════════════════════════════════════════════════════════════════════
add_section_title(doc, "Section 2.  보험부채 잔액")
para = doc.add_paragraph()
run = para.add_run("(단위: 억원,  잔여+발생사고부채 합산)")
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

liab_items = ['BEL', 'RA', 'CSM', '손실부담부채(LOSS)', 'OCI']
tbl2 = doc.add_table(rows=len(liab_items) + 1, cols=4)
tbl2.style = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER

col_widths2 = [Cm(6.0), Cm(4.0), Cm(4.0), Cm(4.0)]
for i, w in enumerate(col_widths2):
    for cell in tbl2.columns[i].cells:
        cell.width = w

hdrs2 = ['항목', '기초(억원)', '기말(억원)', '증감(억원)']
for ci, h in enumerate(hdrs2):
    style_header_cell(tbl2.rows[0].cells[ci], h)

for ri, (item, beg, end, chg) in enumerate(
    zip(liab_items, bal_beg, bal_end, bal_chg), start=1
):
    cells = tbl2.rows[ri].cells
    write_cell(cells[0], item, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    for ci, val in enumerate([beg, end, chg], start=1):
        write_cell(cells[ci], fmt_eok(val) if val is not None else '-')

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# Section 3. 월별 손익 추세 (최근 12개월)
# ══════════════════════════════════════════════════════════════════════════════
add_section_title(doc, "Section 3.  월별 손익 추세 (최근 12개월)")
para = doc.add_paragraph()
run = para.add_run("(단위: 억원)")
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

n_months = len(TREND_COLS)
tbl3 = doc.add_table(rows=4, cols=n_months + 1)
tbl3.style = 'Table Grid'
tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER

# 헤더 행
style_header_cell(tbl3.rows[0].cells[0], '항목', font_size=8)
for ci, hdr in enumerate(trend_headers, start=1):
    style_header_cell(tbl3.rows[0].cells[ci], str(hdr) if hdr else '-',
                      font_size=7)

# 3개 항목
trend_data = [
    ('Ⅰ 보험손익',     trend_ins_pl),
    ('Ⅱ 보험금융손익', trend_fin_pl),
    ('Ⅲ 기타포괄손익', trend_oci),
]
for ri, (name, vals) in enumerate(trend_data, start=1):
    cells = tbl3.rows[ri].cells
    write_cell(cells[0], name, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
               font_size=8)
    for ci, v in enumerate(vals, start=1):
        write_cell(cells[ci], fmt_eok(v, decimals=0) if v is not None else '-',
                   font_size=7)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# Section 4. 예실차 분석 (PL)
# ══════════════════════════════════════════════════════════════════════════════
add_section_title(doc, "Section 4.  예실차 분석")
para = doc.add_paragraph()
run = para.add_run("(단위: 억원,  실제-예상 기준)")
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

tbl4 = doc.add_table(rows=len(var_rows) + 1, cols=4)
tbl4.style = 'Table Grid'
tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER

col_widths4 = [Cm(5.5), Cm(4.0), Cm(4.0), Cm(4.0)]
for i, w in enumerate(col_widths4):
    for cell in tbl4.columns[i].cells:
        cell.width = w

hdrs4 = ['구분', '예상(억원)', '실제(억원)', '예상-실제(억원)']
for ci, h in enumerate(hdrs4):
    style_header_cell(tbl4.rows[0].cells[ci], h)

for ri, (label, exp, act, diff) in enumerate(var_rows, start=1):
    cells = tbl4.rows[ri].cells
    is_total = (label == '계')
    write_cell(cells[0], label, bold=is_total, align=WD_ALIGN_PARAGRAPH.LEFT)
    for ci, val in enumerate([exp, act, diff], start=1):
        write_cell(cells[ci], fmt_eok(val) if val is not None else '-',
                   bold=is_total)
    if is_total:
        for ci in range(4):
            set_cell_bg(cells[ci], 'E2EFDA')

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# Section 5. 이행현금흐름 변동 (가정변경)
# ══════════════════════════════════════════════════════════════════════════════
add_section_title(doc, "Section 5.  이행현금흐름 변동 (가정변경 상세)")
para = doc.add_paragraph()
run = para.add_run("(단위: 원,  각 항목별 변동금액)")
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

for section_name, rows in cf_data.items():
    if not rows:
        continue
    add_subtitle(doc, f"▶ {section_name} 변동")
    tbl_cf = doc.add_table(rows=len(rows) + 1, cols=4)
    tbl_cf.style = 'Table Grid'
    tbl_cf.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_widths_cf = [Cm(6.5), Cm(3.5), Cm(3.5), Cm(3.5)]
    for i, w in enumerate(col_widths_cf):
        for cell in tbl_cf.columns[i].cells:
            cell.width = w

    hdrs_cf = ['변동 구분', '2026.01(원)', '2026.02(원)', '2026.03(원)']
    for ci, h in enumerate(hdrs_cf):
        style_header_cell(tbl_cf.rows[0].cells[ci], h, font_size=9)

    for ri, (code, label, v01, v02, v03) in enumerate(rows, start=1):
        cells = tbl_cf.rows[ri].cells
        write_cell(cells[0], label, align=WD_ALIGN_PARAGRAPH.LEFT)
        for ci, val in enumerate([v01, v02, v03], start=1):
            formatted = f"{int(val):,}" if val is not None and val != 0 else \
                        ('-' if val is None else '0')
            write_cell(cells[ci], formatted)

    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# Section 6. 데이터 미비 항목
# ══════════════════════════════════════════════════════════════════════════════
add_section_title(doc, "Section 6.  데이터 미비 항목 (수동 입력 필요)")

missing_items = [
    ("투자손익",              "별도 투자시스템 데이터 필요"),
    ("간접사업비 배분 명세",  "본사비용 배부 시스템 연계 필요"),
    ("신계약 CSM 상각 계획",  "사업계획 데이터 필요"),
    ("경험조정 RA 조정 근거", "계리사 검토 후 수기 입력"),
    ("자산운용수익률",        "자산운용팀 확인 필요"),
    ("재보험 손익 조정",      "재보험 결산 완료 후 반영"),
    ("공정가치 평가 세부내역","FVPL/FVOCI 구분별 세부 내역"),
]

tbl5 = doc.add_table(rows=len(missing_items) + 1, cols=3)
tbl5.style = 'Table Grid'
tbl5.alignment = WD_TABLE_ALIGNMENT.CENTER

col_widths5 = [Cm(1.5), Cm(5.5), Cm(9.5)]
for i, w in enumerate(col_widths5):
    for cell in tbl5.columns[i].cells:
        cell.width = w

hdrs5 = ['No.', '항목', '비고']
for ci, h in enumerate(hdrs5):
    style_header_cell(tbl5.rows[0].cells[ci], h, bg_color='7F7F7F')

for ri, (item, note) in enumerate(missing_items, start=1):
    cells = tbl5.rows[ri].cells
    write_cell(cells[0], str(ri), align=WD_ALIGN_PARAGRAPH.CENTER)
    write_cell(cells[1], item, align=WD_ALIGN_PARAGRAPH.LEFT)
    write_cell(cells[2], note, align=WD_ALIGN_PARAGRAPH.LEFT)

# 푸터 여백
doc.add_paragraph()
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = para.add_run(f"생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M')}  |  계리결산팀")
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ──────────────────────────────────────────────────────────────────────────────
# 저장
# ──────────────────────────────────────────────────────────────────────────────
doc.save(OUTPUT_FILE)
print(f"\n보고서 생성 완료!")
print(f"출력 경로: {OUTPUT_FILE}")

# ──────────────────────────────────────────────────────────────────────────────
# 데이터 샘플 출력
# ──────────────────────────────────────────────────────────────────────────────
print("\n" + "=" * 60)
print("[ 섹션별 데이터 샘플 ]")
print("=" * 60)
print(f"\n[Section 1 - 손익 요약]")
for _, key, label, _ in pl_rows[:5]:
    cur, mom, cum3 = read_pl_row(key)
    print(f"  {label:<28} | 당월: {fmt_eok(cur):>10} | 전월比: {fmt_eok(mom):>10} | 누계: {fmt_eok(cum3):>10}")

print(f"\n[Section 2 - 보험부채 잔액]")
for item, beg, end, chg in zip(liab_items, bal_beg, bal_end, bal_chg):
    print(f"  {item:<20} | 기초: {fmt_eok(beg):>12} | 기말: {fmt_eok(end):>12} | 증감: {fmt_eok(chg):>12}")

print(f"\n[Section 3 - 최근 12개월 보험손익 (억원)]")
for hdr, val in zip(trend_headers[-6:], trend_ins_pl[-6:]):
    print(f"  {hdr}: {fmt_eok(val)}")

print(f"\n[Section 4 - 예실차]")
for label, exp, act, diff in var_rows:
    print(f"  {label:<14} | 예상: {fmt_eok(exp):>10} | 실제: {fmt_eok(act):>10} | 차이: {fmt_eok(diff):>10}")

print(f"\n[Section 5 - 이행현금흐름 변동 (BEL, 2026.03)]")
for code, label, v01, v02, v03 in cf_data.get('BEL', []):
    if v03 and v03 != 0:
        print(f"  {label:<30}: {int(v03):>18,}")
