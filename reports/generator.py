"""
보고서 자동생성 스크립트
Usage: python generator.py --type [actuary|management|finance] --ym YYYYMM
       --db <actuarial.db> --out <output_dir> [--author <str>]
"""
import argparse, sqlite3, json, sys, os, glob
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("오류: python-docx 미설치. 'pip install python-docx' 실행 후 다시 시도하세요.")
    print(json.dumps({"status":"error","message":"python-docx not installed","filename":""}))
    sys.exit(1)

# ── 헬퍼 ──────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def heading(doc, text, level=1, color='1e3a5f'):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return p

def para(doc, text, bold=False, size=10, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if align:
        p.alignment = align
    return p

def table_header_row(table, headers, bg='1e3a5f', fg='FFFFFF'):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(fg)
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, bg)

def add_table_row(table, values, shade=False):
    row = table.add_row()
    for i, v in enumerate(values):
        cell = row.cells[i]
        cell.text = str(v) if v is not None else '-'
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if shade:
            set_cell_bg(cell, 'f8faff')
    return row

def add_section_box(doc, title, content_lines, bg='eff6ff', border_color='4f86f7'):
    """작은 박스 스타일 섹션."""
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    cell = t.cell(0, 0)
    set_cell_bg(cell, bg.replace('#',''))
    p = cell.paragraphs[0]
    run = p.add_run(f'■ {title}')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(border_color.replace('#',''))
    for line in content_lines:
        cp = cell.add_paragraph(line)
        cp.runs[0].font.size = Pt(9) if cp.runs else Pt(9)
    doc.add_paragraph()

# ── DB 조회 ───────────────────────────────────────────────────────────────────

def query_db(db_path, sql, params=()):
    if not Path(db_path).exists():
        return []
    try:
        conn = sqlite3.connect(db_path)
        conn.row_factory = sqlite3.Row
        rows = conn.execute(sql, params).fetchall()
        conn.close()
        return [dict(r) for r in rows]
    except Exception:
        return []

def get_public_rates(db, ym):
    return query_db(db, "SELECT product_type, rate, note FROM public_interest_rate WHERE ym=? ORDER BY product_type", (ym,))

def get_assumptions(db, ym):
    return query_db(db, "SELECT model_type, assumption_item, value, unit FROM model_assumptions WHERE ym=? ORDER BY model_type, assumption_item", (ym,))

def get_expense_ratios(db, ym):
    return query_db(db, "SELECT product_group, ratio, note FROM expense_ratio WHERE ym=? ORDER BY product_group", (ym,))

# ── 문서 생성 함수들 ───────────────────────────────────────────────────────────

def add_cover(doc, title, ym, author):
    """표지 페이지."""
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('계 리 결 산 팀')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)
    run.bold = True

    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(title)
    run2.font.size = Pt(28)
    run2.bold = True
    run2.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)

    doc.add_paragraph()
    ym_str = f"{ym[:4]}년 {int(ym[4:])}월 기준"
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(ym_str)
    run3.font.size = Pt(16)
    run3.font.color.rgb = RGBColor(0x4f, 0x86, 0xf7)

    doc.add_paragraph()
    doc.add_paragraph()
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run(f'작성일: {datetime.now().strftime("%Y년 %m월 %d일")}')
    run4.font.size = Pt(11)
    run4.font.color.rgb = RGBColor(0x9c, 0xa3, 0xaf)

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run5 = p5.add_run(f'작성자: {author}')
    run5.font.size = Pt(11)
    run5.font.color.rgb = RGBColor(0x9c, 0xa3, 0xaf)
    doc.add_page_break()

def fmt_eok(val, decimals=1):
    if val is None or val == '': return '-'
    try:
        v = float(val)
        return f"{v:,.{decimals}f}"
    except Exception:
        return str(val)


def set_cell_bg2(cell, hex_color):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def style_hdr(cell, text, bg='1F497D', fg='FFFFFF', size=9):
    cell.text = ''
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell.paragraphs[0].add_run(text)
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = RGBColor(int(fg[0:2],16), int(fg[2:4],16), int(fg[4:6],16))
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_bg2(cell, bg)


def write_c(cell, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.RIGHT):
    cell.text = ''
    run = cell.paragraphs[0].add_run(str(text))
    run.font.size = Pt(size)
    run.font.bold = bold
    cell.paragraphs[0].alignment = align
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def sec_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)


def find_input_excel(input_dir, ym):
    """input 폴더에서 기준월이 포함된 엑셀 파일 탐색."""
    patterns = [
        os.path.join(input_dir, f'*{ym}*.xlsx'),
        os.path.join(input_dir, f'*{ym}*.xls'),
        os.path.join(input_dir, '*.xlsx'),
        os.path.join(input_dir, '*.xls'),
    ]
    for pat in patterns:
        files = sorted(glob.glob(pat), key=os.path.getmtime, reverse=True)
        if files:
            return files[0]
    return None


def generate_actuary_from_excel(ym, input_excel, author):
    """엑셀 파일에서 실제 데이터를 읽어 계리보고서 생성."""
    try:
        import openpyxl
        from openpyxl.utils import column_index_from_string
    except ImportError:
        return None, "openpyxl 미설치"

    print(f"처리중: 엑셀 파일 로딩 → {os.path.basename(input_excel)}")
    wb = openpyxl.load_workbook(input_excel, data_only=True)

    # ── Setting_Report: 기준월 ──
    ws_set = wb['Setting_Report']
    report_ym = None
    for r in range(3, 10):
        v = ws_set.cell(row=r, column=3).value
        if v and str(v).isdigit() and len(str(v)) == 6:
            report_ym = str(v); break
    if not report_ym:
        report_ym = ym

    # ── 손익_요약: 당월/전월比/3개월누계 ──
    ws_pl = wb['손익_요약']
    COL_BC = column_index_from_string('BC')
    COL_BD = column_index_from_string('BD')
    COL_BE = column_index_from_string('BE')

    # 실제 행 탐색 (B열 또는 C열 키워드 매칭)
    SEARCH_KEYS = {
        'ins_pl':       ['보험손익', 'Ⅰ 보험손익'],
        'ins_rev':      ['Ⅰ.1', '보험수익'],
        'csm_amort':    ['CSM 상각', 'CSM상각'],
        'ra_change':    ['RA 변동', 'RA변동'],
        'ins_claim':    ['보험금 (예상', '보험금(예상'],
        'acq_cost':     ['신계약비 (예상', '신계약비(예상'],
        'maint_cost':   ['유지비'],
        'inv_mgmt':     ['투자관리비'],
        'loss_adj':     ['손해조사비'],
        'acf_alloc':    ['보험취득CF', '보험취득현금흐름 배분'],
        'lob_alloc':    ['손실부담비용 배분'],
        'ins_svc_exp':  ['Ⅰ.2', '보험서비스비용'],
        'occ_claim':    ['발생보험금'],
        'maint_act':    ['유지비(실제)'],
        'inv_act':      ['투자관리비(실제)'],
        'ladj_act':     ['손해조사비(실제)'],
        'loss_lob':     ['손실부담비용 전'],
        'ins_fin_pl':   ['보험금융손익', 'Ⅱ 보험금융'],
        'ins_fin_rev':  ['Ⅱ.1', '보험금융수익'],
        'ins_fin_exp':  ['Ⅱ.2', '보험금융비용'],
        'oci':          ['기타포괄손익', 'Ⅲ'],
    }

    row_map = {}
    for row in ws_pl.iter_rows(min_row=1, max_row=120, min_col=2, max_col=4):
        for cell in row:
            if cell.value and isinstance(cell.value, str):
                cv = cell.value.strip()
                for key, keywords in SEARCH_KEYS.items():
                    if key not in row_map:
                        if any(k in cv for k in keywords):
                            row_map[key] = cell.row
                            break

    def read_pl(key):
        r = row_map.get(key)
        if not r:
            return None, None, None
        return (ws_pl.cell(row=r, column=COL_BC).value,
                ws_pl.cell(row=r, column=COL_BD).value,
                ws_pl.cell(row=r, column=COL_BE).value)

    # 최근 12개월 추세
    TREND_COLS = list(range(column_index_from_string('AR'), column_index_from_string('BC')+1))
    trend_headers = [ws_pl.cell(row=6, column=c).value for c in TREND_COLS]
    def read_trend(key):
        r = row_map.get(key)
        if not r: return [None]*len(TREND_COLS)
        return [ws_pl.cell(row=r, column=c).value for c in TREND_COLS]

    trend_ins_pl  = read_trend('ins_pl')
    trend_fin_pl  = read_trend('ins_fin_pl')
    trend_oci     = read_trend('oci')

    # 예실차 (BJ~BM 열)
    COL_BJ = column_index_from_string('BJ')
    var_rows = []
    for r in range(5, 20):
        label = ws_pl.cell(row=r, column=COL_BJ).value
        if label and isinstance(label, str) and label.strip():
            exp  = ws_pl.cell(row=r, column=COL_BJ+1).value
            act  = ws_pl.cell(row=r, column=COL_BJ+2).value
            diff = ws_pl.cell(row=r, column=COL_BJ+3).value
            var_rows.append((label.strip(), exp, act, diff))
            if len(var_rows) >= 8: break

    # ── 손익_억: 부채 잔액 ──
    ws_eok = wb['손익_억']
    def read_bal_row(rnum):
        vals = [ws_eok.cell(row=rnum, column=c).value for c in range(5, 10)]
        return vals  # BEL, RA, CSM, LOSS, OCI

    # 잔여보장부채 + 발생사고부채 기말/기초 전체 합산
    def sum_bal_rows(label_kw):
        total = [None] * 5
        for r in range(30, 100):
            lbl = ws_eok.cell(row=r, column=4).value
            if lbl and label_kw in str(lbl):
                v = read_bal_row(r)
                if any(x for x in v if x):
                    for i in range(5):
                        if v[i] is not None:
                            total[i] = (total[i] or 0) + v[i]
        return total

    bal_end = sum_bal_rows('기말')
    bal_beg = sum_bal_rows('기초')

    def safe_sub(a, b):
        try: return (a or 0) - (b or 0)
        except: return None

    bal_chg = [safe_sub(bal_end[i], bal_beg[i]) for i in range(5)]

    # ── 이행현금흐름상세 ──
    ws_cf = wb['이행현금흐름상세']
    cf_labels = {
        '006': '가정변경 (할인율/공시이율)', '151': '모델변경',
        '152': '계리가정변경(사업비율)', '153': '계리가정변경(위험률)',
        '154': '계리가정변경(해지율)', '155': '계리가정변경(기타)',
        '185': 'VFA 기업의 몫 조정',
    }
    cf_rows = []
    for row in ws_cf.iter_rows(min_row=2, max_row=80):
        code = row[0].value
        if code and str(code).strip() in cf_labels:
            lbl = cf_labels[str(code).strip()]
            vals = [c.value for c in row[1:4]]
            cf_rows.append((lbl, vals))

    # ── Word 문서 생성 ──
    print(f"처리중: Word 문서 구성 중 (기준월: {report_ym})")
    doc = Document()
    from docx.enum.table import WD_TABLE_ALIGNMENT

    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)
    try:
        from docx.oxml.ns import qn as _qn
        style._element.rPr.rFonts.set(_qn('w:eastAsia'), '맑은 고딕')
    except Exception: pass

    for sec in doc.sections:
        sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.0)

    ym_yr = report_ym[:4]; ym_mo = int(report_ym[4:])

    # 표지
    add_cover(doc, 'IFRS17 계리결산 보고서', report_ym, author)

    # Section 1. 손익 요약
    sec_title(doc, 'Section 1.  손익 요약')
    _p = doc.add_paragraph(); _p.add_run('(단위: 억원)').font.size = Pt(8)

    pl_rows_def = [
        (False,'ins_pl',     'Ⅰ  보험손익',          True),
        (False,'ins_rev',    'Ⅰ.1  보험수익',         True),
        (True, 'csm_amort',  '  - CSM 상각',          False),
        (True, 'ra_change',  '  - RA 변동',           False),
        (True, 'ins_claim',  '  - 보험금(예상-실제)', False),
        (True, 'acq_cost',   '  - 신계약비(예상-실제)',False),
        (True, 'maint_cost', '  - 유지비',            False),
        (True, 'inv_mgmt',   '  - 투자관리비',        False),
        (True, 'loss_adj',   '  - 손해조사비',        False),
        (True, 'acf_alloc',  '  - 보험취득CF배분',    False),
        (True, 'lob_alloc',  '  - 손실부담비용배분',  False),
        (False,'ins_svc_exp','Ⅰ.2  보험서비스비용',  True),
        (True, 'occ_claim',  '  - 발생보험금',        False),
        (True, 'maint_act',  '  - 유지비(실제)',      False),
        (True, 'inv_act',    '  - 투자관리비(실제)',  False),
        (True, 'ladj_act',   '  - 손해조사비(실제)',  False),
        (True, 'loss_lob',   '  - 손실부담비용전/환입',False),
        (False,'ins_fin_pl', 'Ⅱ  보험금융손익',      True),
        (False,'ins_fin_rev','Ⅱ.1  보험금융수익',    True),
        (False,'ins_fin_exp','Ⅱ.2  보험금융비용',    True),
        (False,'oci',        'Ⅲ  기타포괄손익',      True),
    ]

    t1 = doc.add_table(rows=len(pl_rows_def)+1, cols=4)
    t1.style = 'Table Grid'; t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, w in enumerate([Cm(7.0), Cm(3.5), Cm(3.5), Cm(4.0)]):
        for cell in t1.columns[i].cells: cell.width = w
    for ci, (h, bg) in enumerate(zip(['항목','당월(억원)','전월比(억원)','3개월누계(억원)'],
                                     ['1F497D','2E74B5','2E74B5','2E74B5'])):
        style_hdr(t1.rows[0].cells[ci], h, bg=bg)

    for ri, (_, key, label, bold) in enumerate(pl_rows_def, start=1):
        cur, mom, cum3 = read_pl(key)
        cells = t1.rows[ri].cells
        write_c(cells[0], label, bold=bold, align=WD_ALIGN_PARAGRAPH.LEFT)
        for ci, val in enumerate([cur, mom, cum3], start=1):
            write_c(cells[ci], fmt_eok(val) if val is not None else '-', bold=bold)
        if bold:
            for ci in range(4): set_cell_bg2(cells[ci], 'DDEEFF')

    doc.add_paragraph()

    # Section 2. 보험부채 잔액
    sec_title(doc, 'Section 2.  보험부채 잔액')
    doc.add_paragraph().add_run('(단위: 억원)').font.size = Pt(8)

    liab_names = ['BEL', 'RA', 'CSM', '손실부담부채(LOSS)', 'OCI']
    t2 = doc.add_table(rows=6, cols=4)
    t2.style = 'Table Grid'; t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, w in enumerate([Cm(6.0), Cm(4.0), Cm(4.0), Cm(4.0)]):
        for cell in t2.columns[i].cells: cell.width = w
    for ci, h in enumerate(['항목','기초(억원)','기말(억원)','증감(억원)']):
        style_hdr(t2.rows[0].cells[ci], h)
    for ri, (name, beg, end, chg) in enumerate(zip(liab_names, bal_beg, bal_end, bal_chg), start=1):
        cells = t2.rows[ri].cells
        write_c(cells[0], name, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
        for ci, v in enumerate([beg, end, chg], start=1):
            write_c(cells[ci], fmt_eok(v) if v is not None else '-')

    doc.add_paragraph()

    # Section 3. 월별 추세
    sec_title(doc, 'Section 3.  월별 손익 추세 (최근 12개월)')
    doc.add_paragraph().add_run('(단위: 억원)').font.size = Pt(8)

    n_cols = len(TREND_COLS)
    t3 = doc.add_table(rows=4, cols=n_cols+1)
    t3.style = 'Table Grid'; t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    style_hdr(t3.rows[0].cells[0], '항목', size=8)
    for ci, hdr in enumerate(trend_headers, start=1):
        style_hdr(t3.rows[0].cells[ci], str(hdr) if hdr else '-', size=7)
    for ri, (name, vals) in enumerate([
        ('Ⅰ 보험손익', trend_ins_pl), ('Ⅱ 보험금융손익', trend_fin_pl), ('Ⅲ 기타포괄손익', trend_oci)
    ], start=1):
        cells = t3.rows[ri].cells
        write_c(cells[0], name, bold=True, size=8, align=WD_ALIGN_PARAGRAPH.LEFT)
        for ci, v in enumerate(vals, start=1):
            write_c(cells[ci], fmt_eok(v, 0) if v is not None else '-', size=7)

    doc.add_paragraph()

    # Section 4. 예실차
    if var_rows:
        sec_title(doc, 'Section 4.  예실차 분석')
        doc.add_paragraph().add_run('(단위: 억원)').font.size = Pt(8)
        t4 = doc.add_table(rows=len(var_rows)+1, cols=4)
        t4.style = 'Table Grid'; t4.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ci, h in enumerate(['구분','예상(억원)','실제(억원)','예상-실제(억원)']):
            style_hdr(t4.rows[0].cells[ci], h)
        for ri, (lbl, exp, act, diff) in enumerate(var_rows, start=1):
            cells = t4.rows[ri].cells
            is_tot = '계' in lbl
            write_c(cells[0], lbl, bold=is_tot, align=WD_ALIGN_PARAGRAPH.LEFT)
            for ci, v in enumerate([exp, act, diff], start=1):
                write_c(cells[ci], fmt_eok(v) if v is not None else '-', bold=is_tot)
            if is_tot:
                for ci in range(4): set_cell_bg2(cells[ci], 'E2EFDA')
        doc.add_paragraph()

    # Section 5. 이행현금흐름 변동
    if cf_rows:
        sec_title(doc, 'Section 5.  이행현금흐름 변동 (가정변경 상세)')
        doc.add_paragraph().add_run('(단위: 원)').font.size = Pt(8)
        t5 = doc.add_table(rows=len(cf_rows)+1, cols=4)
        t5.style = 'Table Grid'; t5.alignment = WD_TABLE_ALIGNMENT.CENTER
        prev2 = f"{ym_yr}{str(ym_mo-2).zfill(2)}" if ym_mo > 2 else f"{int(ym_yr)-1}{str(12+ym_mo-2).zfill(2)}"
        prev1 = f"{ym_yr}{str(ym_mo-1).zfill(2)}" if ym_mo > 1 else f"{int(ym_yr)-1}12"
        for ci, h in enumerate(['변동 구분', f'{prev2[:4]}.{prev2[4:]}(원)', f'{prev1[:4]}.{prev1[4:]}(원)', f'{ym_yr}.{str(ym_mo).zfill(2)}(원)']):
            style_hdr(t5.rows[0].cells[ci], h, size=9)
        for ri, (lbl, vals) in enumerate(cf_rows, start=1):
            cells = t5.rows[ri].cells
            write_c(cells[0], lbl, align=WD_ALIGN_PARAGRAPH.LEFT)
            for ci, v in enumerate(vals, start=1):
                txt = f"{int(v):,}" if v and v != 0 else ('-' if v is None else '0')
                write_c(cells[ci], txt)
        doc.add_paragraph()

    # Section 6. 데이터 미비 항목
    sec_title(doc, 'Section 6.  데이터 미비 항목 (수동 입력 필요)')
    t6 = doc.add_table(rows=8, cols=3)
    t6.style = 'Table Grid'; t6.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, h in enumerate(['No.', '항목', '비고']):
        style_hdr(t6.rows[0].cells[ci], h, bg='7F7F7F')
    for ri, (item, note) in enumerate([
        ('투자손익',              '별도 투자시스템 데이터 필요'),
        ('간접사업비 배분 명세',  '본사비용 배부 시스템 연계 필요'),
        ('신계약 CSM 상각 계획',  '사업계획 데이터 필요'),
        ('경험조정 RA 조정 근거', '계리사 검토 후 수기 입력'),
        ('자산운용수익률',        '자산운용팀 확인 필요'),
        ('재보험 손익 조정',      '재보험 결산 완료 후 반영'),
        ('공정가치 평가 세부내역','FVPL/FVOCI 구분별 세부 내역'),
    ], start=1):
        cells = t6.rows[ri].cells
        write_c(cells[0], str(ri), align=WD_ALIGN_PARAGRAPH.CENTER)
        write_c(cells[1], item, align=WD_ALIGN_PARAGRAPH.LEFT)
        write_c(cells[2], note, align=WD_ALIGN_PARAGRAPH.LEFT)

    # 푸터
    doc.add_paragraph()
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = fp.add_run(f"생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M')}  |  {author}")
    run.font.size = Pt(8); run.font.color.rgb = RGBColor(0x80,0x80,0x80)

    return doc, None


def generate_actuary(ym, db_path, author):
    """계리보고서 생성 - 엑셀 파일이 있으면 실제 데이터 사용, 없으면 플레이스홀더."""
    input_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'input')
    excel_file = find_input_excel(input_dir, ym)

    if excel_file:
        print(f"처리중: 엑셀 데이터 감지됨 → {os.path.basename(excel_file)}")
        doc, err = generate_actuary_from_excel(ym, excel_file, author)
        if doc:
            return doc
        print(f"처리중: 엑셀 읽기 실패({err}), 플레이스홀더로 생성")
    else:
        print(f"처리중: input 폴더에 {ym} 엑셀 파일 없음 → 플레이스홀더로 생성")

    # ── 플레이스홀더 버전 ──
    print(f"처리중: 계리보고서 작성 중 (기준월: {ym})")
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)
    for section in doc.sections:
        section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3);  section.right_margin  = Cm(2.5)

    add_cover(doc, '계리보고서', ym, author)
    ym_str = f"{ym[:4]}년 {int(ym[4:])}월"

    heading(doc, f'1. {ym_str} 결산 개요')
    para(doc, f'본 보고서는 {ym_str} 기준 계리결산 결과를 요약한 것입니다. '
         f'IFRS17 기준에 따라 보험계약마진(CSM), 최선추정부채(BEL), 위험조정(RA) 등 주요 지표를 포함합니다.')
    doc.add_paragraph()

    add_section_box(doc, '주요 결산 일정', [
        f'· 결산 기준일: {ym_str} 말일',
        '· 초안 작성: 결산월 +5 영업일',
        '· 내부 검토: 결산월 +10 영업일',
        '· 최종 제출: 결산월 +15 영업일',
    ])

    heading(doc, '2. 주요 재무 수치')
    para(doc, f'※ reports/input/ 폴더에 {ym} 엑셀 파일을 넣으면 실제 데이터로 자동 생성됩니다.',
         color='9ca3af', size=9)
    doc.add_paragraph()

    t = doc.add_table(rows=1, cols=4)
    t.style = 'Table Grid'
    table_header_row(t, ['구분', '금 액(억원)', '전월비', '비고'])
    for row_data, shade in [
        (['잔여보장부채 (BEL)', '-', '-', '최선추정부채'],  False),
        (['잔여보장부채 (RA)',  '-', '-', '위험조정'],       True),
        (['잔여보장부채 (CSM)', '-', '-', '보험계약마진'],   False),
        (['발생사고부채',       '-', '-', ''],                True),
        (['합 계',             '-', '-', ''],                False),
    ]:
        add_table_row(t, row_data, shade)
    doc.add_paragraph()

    heading(doc, '3. 공시이율 현황')
    rates = get_public_rates(db_path, ym)
    if rates:
        t2 = doc.add_table(rows=1, cols=3); t2.style = 'Table Grid'
        table_header_row(t2, ['상품유형', '공시이율(%)', '비고'])
        for i, r in enumerate(rates):
            add_table_row(t2, [r['product_type'], f"{r['rate']:.2f}%", r.get('note','')], i%2==1)
    else:
        para(doc, f'※ {ym_str} 공시이율 데이터 없음.', color='9ca3af', size=9)
    doc.add_paragraph()

    heading(doc, '4. 특이사항 및 의견')
    para(doc, '· (특이사항 없음 또는 직접 입력)')
    doc.add_paragraph()

    return doc

def generate_management(ym, db_path, author):
    """경영개선협의회 보고서 생성."""
    print(f"처리중: 경영개선협의회 보고서 작성 중 (기준월: {ym})")
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)
    for section in doc.sections:
        section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3); section.right_margin = Cm(2.5)

    add_cover(doc, '경영개선협의회\n계리결산 현황 보고', ym, author)
    ym_str = f"{ym[:4]}년 {int(ym[4:])}월"

    heading(doc, '1. 보고 배경 및 목적')
    para(doc, f'· {ym_str} 기준 계리결산 주요 현황 및 이슈사항을 경영개선협의회에 보고합니다.')
    para(doc, '· IFRS17 시행에 따른 결산 프로세스 안정화 현황 및 개선 과제를 공유합니다.')
    doc.add_paragraph()

    heading(doc, '2. 주요 성과 지표 (KPI)')
    t = doc.add_table(rows=1, cols=4)
    t.style = 'Table Grid'
    table_header_row(t, ['지표명', '목표', '실적', '달성률'])
    for row_data, shade in [
        (['결산 기한 준수율', '100%', '-', '-'],       False),
        (['결산 오류 건수', '0건', '-', '-'],           True),
        (['자동화율', '80%', '-', '-'],                 False),
        (['보고서 적시 제출', '100%', '-', '-'],        True),
    ]:
        add_table_row(t, row_data, shade)
    doc.add_paragraph()

    heading(doc, '3. 계리결산 주요 결과')
    rates = get_public_rates(db_path, ym)
    if rates:
        para(doc, f'■ {ym_str} 공시이율 현황', bold=True)
        t2 = doc.add_table(rows=1, cols=3)
        t2.style = 'Table Grid'
        table_header_row(t2, ['상품유형', '공시이율(%)', '비고'])
        for i, r in enumerate(rates):
            add_table_row(t2, [r['product_type'], f"{r['rate']:.2f}%", r.get('note','')], i%2==1)
    else:
        add_section_box(doc, f'{ym_str} 결산 결과 요약', [
            '· 잔여보장부채 합계: (직접 입력)',
            '· 발생사고부채: (직접 입력)',
            '· CSM 변동 주요 원인: (직접 입력)',
        ])
    doc.add_paragraph()

    heading(doc, '4. 개선 과제 및 대응 방안')
    t3 = doc.add_table(rows=1, cols=4)
    t3.style = 'Table Grid'
    table_header_row(t3, ['과제', '현황', '대응 방안', '완료 예정'])
    for row_data, shade in [
        (['결산 자동화', '진행중', '파이썬 스크립트 개발', '2026.06'],   False),
        (['검증 프로세스 고도화', '계획중', '체크리스트 표준화', '2026.09'],  True),
        (['직접 입력', '-', '-', '-'],                                    False),
    ]:
        add_table_row(t3, row_data, shade)
    doc.add_paragraph()

    heading(doc, '5. 향후 일정')
    para(doc, f'· 차기 결산 기준일: {ym[:4]}년 {int(ym[4:])+1 if int(ym[4:]) < 12 else 1}월 말')
    para(doc, '· 차기 보고 일정: 결산 완료 후 +15 영업일')
    para(doc, '· 기타: (직접 입력)')

    return doc

def generate_finance(ym, db_path, author):
    """재무협의체 보고서 생성."""
    print(f"처리중: 재무협의체 보고서 작성 중 (기준월: {ym})")
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)
    for section in doc.sections:
        section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3); section.right_margin = Cm(2.5)

    add_cover(doc, '재무협의체\nIFRS17·K-ICS 현황 보고', ym, author)
    ym_str = f"{ym[:4]}년 {int(ym[4:])}월"

    heading(doc, '1. 재무 현황 요약')
    add_section_box(doc, f'{ym_str} 재무 현황', [
        '· 보험계약마진(CSM): (직접 입력) 억원',
        '· 최선추정부채(BEL): (직접 입력) 억원',
        '· 위험조정(RA): (직접 입력) 억원',
        '· 전월 대비 주요 변동: (직접 입력)',
    ])

    heading(doc, '2. IFRS17 주요 지표')
    t = doc.add_table(rows=1, cols=5)
    t.style = 'Table Grid'
    table_header_row(t, ['지표', '당월', '전월', '전월비', '비고'])
    for row_data, shade in [
        (['CSM', '-', '-', '-', ''],               False),
        (['BEL (잔여보장)', '-', '-', '-', ''],    True),
        (['RA', '-', '-', '-', ''],                 False),
        (['발생사고부채', '-', '-', '-', ''],       True),
        (['보험수익', '-', '-', '-', ''],            False),
    ]:
        add_table_row(t, row_data, shade)
    doc.add_paragraph()

    heading(doc, '3. K-ICS 현황')
    t2 = doc.add_table(rows=1, cols=4)
    t2.style = 'Table Grid'
    table_header_row(t2, ['구분', '당월', '전분기', '비고'])
    for row_data, shade in [
        (['K-ICS 비율', '-', '-', '규제 기준 100% 이상'],  False),
        (['가용자본', '-', '-', ''],                        True),
        (['요구자본', '-', '-', ''],                        False),
    ]:
        add_table_row(t2, row_data, shade)
    doc.add_paragraph()

    heading(doc, '4. 리스크 모니터링')
    assumptions = get_assumptions(db_path, ym)
    if assumptions:
        para(doc, f'■ {ym_str} 주요 가정 현황', bold=True)
        t3 = doc.add_table(rows=1, cols=4)
        t3.style = 'Table Grid'
        table_header_row(t3, ['모형', '가정항목', '값', '단위'])
        for i, a in enumerate(assumptions):
            add_table_row(t3, [a['model_type'], a['assumption_item'], a['value'], a.get('unit','')], i%2==1)
    else:
        add_section_box(doc, '주요 리스크 항목', [
            '· 금리 리스크: (직접 입력)',
            '· 보험 리스크: (직접 입력)',
            '· 신용 리스크: (직접 입력)',
        ])
    doc.add_paragraph()

    heading(doc, '5. 차기 예정 사항')
    para(doc, f'· 차기 기준월: {ym[:4]}년 {int(ym[4:])+1 if int(ym[4:]) < 12 else 1}월')
    para(doc, '· 주요 변경 예정 사항: (직접 입력)')
    para(doc, '· 특이사항: (직접 입력)')

    return doc

# ── 메인 ──────────────────────────────────────────────────────────────────────

GENERATORS = {
    'actuary':    (generate_actuary,    '계리보고서'),
    'management': (generate_management, '경영개선협의회보고서'),
    'finance':    (generate_finance,    '재무협의체보고서'),
}

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--type', required=True, choices=GENERATORS.keys())
    ap.add_argument('--ym', required=True)
    ap.add_argument('--db', required=True)
    ap.add_argument('--out', required=True)
    ap.add_argument('--author', default='계리결산팀')
    args = ap.parse_args()

    gen_fn, label = GENERATORS[args.type]
    print(f"처리중: {label} 생성 시작 (기준월: {args.ym})")
    print(f"처리중: DB 경로 확인 → {args.db}")

    doc = gen_fn(args.ym, args.db, args.author)

    Path(args.out).mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"{args.type}_{args.ym}_{ts}.docx"
    out_path = Path(args.out) / filename
    doc.save(str(out_path))

    print(f"생성 완료: {filename}")
    print(json.dumps({
        "status": "success",
        "filename": filename,
        "path": str(out_path),
        "message": f"{label} 생성 완료"
    }, ensure_ascii=False))

if __name__ == '__main__':
    try:
        main()
    except Exception as e:
        print(f"오류: {e}")
        print(json.dumps({"status":"error","message":str(e),"filename":""}, ensure_ascii=False))
        sys.exit(1)
